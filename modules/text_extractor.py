"""
Extract text from PDF and DOCX files

Copyright (c) 2025 Mattias Nyqvist
Licensed under the MIT License
"""

import PyPDF2
import pdfplumber
from docx import Document
from typing import Tuple, Optional


def extract_text_from_pdf(file_path: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract text from PDF file.
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Tuple of (extracted_text, error_message)
    """
    try:
        # Try pdfplumber first (better for complex PDFs)
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        
        if text.strip():
            return text, None
        
        # Fallback to PyPDF2
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        
        if not text.strip():
            return None, "Could not extract text from PDF. The file might be scanned or image-based."
        
        return text, None
        
    except Exception as e:
        return None, f"Error reading PDF: {str(e)}"


def extract_text_from_docx(file_path: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract text from DOCX file.
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Tuple of (extracted_text, error_message)
    """
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n\n"
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += row_text + "\n"
            text += "\n"
        
        if not text.strip():
            return None, "Could not extract text from document. The file appears to be empty."
        
        return text, None
        
    except Exception as e:
        return None, f"Error reading DOCX: {str(e)}"


def extract_text(file_path: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract text from file (auto-detect PDF or DOCX).
    
    Args:
        file_path: Path to file
        
    Returns:
        Tuple of (extracted_text, error_message)
    """
    file_ext = file_path.lower().split('.')[-1]
    
    if file_ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext in ['docx', 'doc']:
        return extract_text_from_docx(file_path)
    else:
        return None, f"Unsupported file type: .{file_ext}"


def clean_text(text: str) -> str:
    """
    Clean extracted text.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    # Remove excessive whitespace
    lines = [line.strip() for line in text.split('\n')]
    lines = [line for line in lines if line]
    
    # Join with single newline
    cleaned = '\n'.join(lines)
    
    # Remove multiple spaces
    import re
    cleaned = re.sub(r' +', ' ', cleaned)
    
    return cleaned